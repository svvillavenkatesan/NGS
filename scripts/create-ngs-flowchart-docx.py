from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'NGS_Current_Flowchart_A3.docx'
ASSETS = ROOT / 'tmp' / 'ngs-flowchart-assets'
ASSETS.mkdir(parents=True, exist_ok=True)
OUT.parent.mkdir(parents=True, exist_ok=True)

W, H = 3308, 2338
NAVY = '#102A43'; BLUE = '#1967A3'; TEAL = '#087F8C'; GREEN = '#2E7D32'
GOLD = '#C58B12'; RED = '#B3261E'; LIGHT = '#F5F8FC'; INK = '#132238'
MUTED = '#52677D'; LINE = '#7790A8'; WHITE = '#FFFFFF'
FONT = r'C:\Windows\Fonts\Nirmala.ttc'
FONT_BOLD = r'C:\Windows\Fonts\Nirmala.ttc'

def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT, size)

def center_text(draw, box, text, size=34, color=INK, bold=False, spacing=8):
    f = font(size, bold)
    lines = text.split('\n')
    heights = [draw.textbbox((0,0), line, font=f)[3] for line in lines]
    total = sum(heights) + spacing * (len(lines)-1)
    y = box[1] + (box[3]-box[1]-total)/2
    for line, hh in zip(lines, heights):
        bb = draw.textbbox((0,0), line, font=f)
        x = box[0] + (box[2]-box[0]-(bb[2]-bb[0]))/2
        draw.text((x,y), line, font=f, fill=color)
        y += hh + spacing

def box(draw, xy, title, subtitle='', fill='#EAF2F8', outline=BLUE, title_size=37, subtitle_size=27):
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=5)
    if subtitle:
        mid = xy[1] + int((xy[3]-xy[1])*.43)
        center_text(draw, (xy[0]+12,xy[1]+5,xy[2]-12,mid), title, title_size, NAVY, True)
        center_text(draw, (xy[0]+16,mid,xy[2]-16,xy[3]-8), subtitle, subtitle_size, MUTED, False, 5)
    else:
        center_text(draw, xy, title, title_size, NAVY, True)

def arrow(draw, start, end, color=LINE, width=8):
    draw.line([start,end], fill=color, width=width)
    import math
    ang = math.atan2(end[1]-start[1], end[0]-start[0])
    length = 24
    p1 = (end[0]-length*math.cos(ang-0.55), end[1]-length*math.sin(ang-0.55))
    p2 = (end[0]-length*math.cos(ang+0.55), end[1]-length*math.sin(ang+0.55))
    draw.polygon([end,p1,p2], fill=color)

def header(draw, page_title, page_no):
    draw.rectangle((0,0,W,150), fill=NAVY)
    draw.text((80,42), 'NGS · CURRENT SYSTEM FLOWCHART', font=font(50, True), fill=WHITE)
    r = draw.textbbox((0,0), page_title, font=font(35, True))
    draw.text((W-80-(r[2]-r[0]),54), page_title, font=font(35, True), fill='#9FD8FF')
    draw.text((80,H-72), 'Current architecture · 31 August 2026', font=font(25), fill=MUTED)
    draw.text((W-180,H-72), f'Page {page_no}/2', font=font(25, True), fill=MUTED)

def page1():
    im = Image.new('RGB',(W,H),WHITE); d=ImageDraw.Draw(im); header(d,'MANAGEMENT & CONTROL',1)
    owner=(1250,215,2058,420); box(d,owner,'SYSTEM OWNER','Secure owner login · Full control', '#FFF6D9', GOLD)
    a1=(250,590,1020,820); a2=(1269,590,2039,820); a3=(2288,590,3058,820)
    for i,b in enumerate([a1,a2,a3],1): box(d,b,f'SUPER ADMIN {i}',f'ID: YYMM#### · Individual seller limit\nOwn passwords · Own seller reports','#E8F2FB',BLUE)
    arrow(d,(1654,420),(635,590)); arrow(d,(1654,420),(1654,590)); arrow(d,(1654,420),(2673,590))
    sellers=[(180,1040,1090,1320),(1199,1040,2109,1320),(2218,1040,3128,1320)]
    for i,b in enumerate(sellers,1):
        box(d,b,'DIRECT SELLERS',f'Created / Active / Limit / Remaining\nSeller IDs belong only to Super Admin {i}','#EAF7EE',GREEN)
        arrow(d,((a1,a2,a3)[i-1][0]+385,820),(b[0]+455,1040))
    master=(200,1545,1525,1840); reports=(1783,1545,3108,1840)
    box(d,master,'SYSTEM-WIDE MASTER SETTINGS','Lot Codes: KL / DR · Show timings\nScheme master · Minimum rate · MRP · Prize structure','#F1EDFA','#6941C6')
    box(d,reports,'OWNER CONSOLIDATED REPORT','Each Super Admin separately\nDaily · Weekly · Monthly · Sales · Prize · Bonus · Profit/Loss','#FFF1E8','#C35A15')
    for b in [a1,a2,a3]: arrow(d,(b[0]+385,820),(862,1545), '#6941C6',6)
    for b in sellers: arrow(d,(b[0]+455,1320),(2445,1545), '#C35A15',6)
    note=(480,1980,2828,2200); box(d,note,'CURRENT ISOLATION RULE','Seller accounts, limits, reports and action passwords are separate per Super Admin.\nLot Code, scheme master and result scope are currently shared system-wide.','#FDECEC',RED,34,27)
    return im

def page2():
    im=Image.new('RGB',(W,H),WHITE); d=ImageDraw.Draw(im); header(d,'ENTRY · RESULT · REPORT · DEPLOYMENT',2)
    y=235; gap=55; bw=535; bh=205; xs=[70,660,1250,1840,2430]
    top=[('ANDROID SELLER','Assigned Lot / Scheme'),('SHOW CLOCK','Date · Time · Closing countdown'),('TICKET ENTRY','Number · BOX · Qty · Quick Qty'),('CURRENT BILL','Latest entry first · Total'),('COMPLETE','OK / OK & Print')]
    for x,(t,s) in zip(xs,top): box(d,(x,y,x+bw,y+bh),t,s,'#EAF7EE',GREEN,34,25)
    for i in range(4): arrow(d,(xs[i]+bw,y+bh//2),(xs[i+1],y+bh//2))
    vy=600; validations=[('ACTIVE?','Seller account'),('ASSIGNED?','Lot code & scheme'),('OPEN?','Show closing time'),('VALID?','Rate & quantity'),('ACCEPT','Save ticket/bill')]
    for x,(t,s) in zip(xs,validations): box(d,(x,vy,x+bw,vy+bh),t,s,'#E8F2FB',BLUE,34,25)
    arrow(d,(337,440),(337,600))
    for i in range(4): arrow(d,(xs[i]+bw,vy+bh//2),(xs[i+1],vy+bh//2))
    d.text((85,842),'Any failed validation → Entry rejected',font=font(28,True),fill=RED)
    ry=955; result=[('ENTRY CLOSES','Configured show time'),('WAIT 1 MINUTE','Publishing blocked'),('RESULT PANEL','Date · Lot · Show · 4 digits'),('PROFIT PREVIEW','Sales · Prize exposure · P/L'),('PASSWORD','Confirm publish')]
    for x,(t,s) in zip(xs,result): box(d,(x,ry,x+bw,ry+bh),t,s,'#FFF6D9',GOLD,32,24)
    for i in range(4): arrow(d,(xs[i]+bw,ry+bh//2),(xs[i+1],ry+bh//2))
    ly=1315; lock=(250,ly,970,ly+245); settle=(1294,ly,2014,ly+245); live=(2338,ly,3058,ly+245)
    box(d,lock,'PERMANENT LOCK','Lot Code + Show + Date\nCannot edit or republish','#FDECEC',RED,33,25)
    box(d,settle,'SETTLEMENT','Every matching ticket\nWIN / LOSE · Prize','#EAF7EE',GREEN,33,25)
    box(d,live,'LIVE RESULT','Visible to Owner\nSuper Admin · Sellers','#E8F2FB',BLUE,33,25)
    arrow(d,(2697,1160),(610,1315),GOLD,7); arrow(d,(970,1437),(1294,1437)); arrow(d,(2014,1437),(2338,1437))
    by=1770; blocks=[(90,by,790,2115,'SELLER REPORT','Entry · Bill · Item\nWinning · Payment · A4'),(865,by,1565,2115,'ADMIN REPORT','Own sellers only\nDaily · Weekly · Monthly'),(1640,by,2340,2115,'STORAGE','Node API · PostgreSQL\nAudit log · Verified backup'),(2415,by,3115,2115,'DEPLOYMENT','Domain · HTTPS · Caddy\nSigned APK · Same key updates')]
    for x1,y1,x2,y2,t,s in blocks: box(d,(x1,y1,x2,y2),t,s,'#F5F8FC',TEAL,32,24)
    arrow(d,(1654,1560),(440,1770),GREEN,6); arrow(d,(1654,1560),(1215,1770),GREEN,6); arrow(d,(2698,1560),(1990,1770),BLUE,6); arrow(d,(2698,1560),(2765,1770),BLUE,6)
    # Word can crop the first pixels of an image that starts on a forced new page.
    # Keep a small white safety margin above page 2 and scale the artwork to fit.
    padded = Image.new('RGB', (W, H), WHITE)
    padded.paste(im.resize((W, H-90), Image.Resampling.LANCZOS), (0, 90))
    return padded

images=[]
for idx, maker in enumerate((page1,page2),1):
    path=ASSETS/f'page-{idx}.png'; maker().save(path, quality=96); images.append(path)

doc=Document(); sec=doc.sections[0]
sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width=Inches(16.54); sec.page_height=Inches(11.69)
sec.top_margin=Inches(.35); sec.bottom_margin=Inches(.35); sec.left_margin=Inches(.45); sec.right_margin=Inches(.45)
sec.header_distance=Inches(.15); sec.footer_distance=Inches(.15)
style=doc.styles['Normal']; style.font.name='Nirmala UI'; style.font.size=Pt(10)
style._element.rPr.rFonts.set(qn('w:ascii'),'Nirmala UI'); style._element.rPr.rFonts.set(qn('w:hAnsi'),'Nirmala UI')
for idx,path in enumerate(images):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
    if idx > 0: p.paragraph_format.page_break_before = True
    p.add_run().add_picture(str(path), width=Inches(15.0), height=Inches(10.3))

props=doc.core_properties; props.title='NGS Current System Flowchart - A3'; props.subject='Owner, Super Admin, Seller, Entry, Result, Reports, Security and Deployment'; props.author='SVV'; props.keywords='NGS, flowchart, A3, system architecture'
doc.save(OUT)
print(OUT)
